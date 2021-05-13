from random import choice
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
import xlrd
import pandas as pd
import time,datetime
from docx.oxml.ns import qn
from docx.shared import RGBColor,Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn



#获取时间
todaytime=datetime.datetime.now()
todaytime1 = str(todaytime.year)+todaytime.strftime('%m')+ todaytime.strftime('%d')
todaytime2 = str(todaytime.year)+'-'+todaytime.strftime('%m')+'-'+ todaytime.strftime('%d')
todaytime3 = todaytime.strftime('%m')+'-'+ todaytime.strftime('%d')

#创建表格
df1 = pd.read_excel('基础数据整理-0423.xlsx',skiprows=38)
df1 = pd.DataFrame(df1)
a = df1.iloc[:16,:9]
b = df1.iloc[:16,10:15]
b['日期'] =todaytime2
b.rename(columns={'地市.1': '地市'}, inplace=True)
b=b[['地市','日期','上月用户数','日开机用户数','日收视用户数','收视率']]
b['收视率']=b['收视率'].apply(lambda x: '%.2f%%' % (x * 100))  #转换百分比
a['有线接入率']=a['有线接入率'].apply(lambda x: '%.2f%%' % (x * 100))  #转换百分比
a['EPG响应成功率']=a['EPG响应成功率'].apply(lambda x: '%.2f%%' % (x * 100))  #转换百分比
a['播放成功率']=a['播放成功率'].apply(lambda x: '%.2f%%' % (x * 100))  #转换百分比
a['卡顿时长占比']=a['卡顿时长占比'].apply(lambda x: '%.4f%%' % (x * 100))  #转换百分比
a['卡顿用户占比']=a['卡顿用户占比'].apply(lambda x: '%.2f%%' % (x * 100))  #转换百分比
a['视频播放优良率d-OA']=a['视频播放优良率d-OA'].apply(lambda x: '%.2f%%' % (x * 100))  #转换百分比

#提取需要数据
data1 = b.loc[a.shape[0]-1,'上月用户数']
data2 = b.loc[a.shape[0]-1,'日收视用户数']
data3 = b.loc[a.shape[0]-1,'收视率']

with pd.ExcelWriter('广西机顶盒软探针指标分析报告-'+todaytime3 + '.xlsx') as writer:  # 写入结果为当前路径
    b.to_excel(writer, sheet_name='1', startcol=0, startrow=0, index=False, header=True)
    a.to_excel(writer, sheet_name='2', startcol=0, startrow=0, index=False, header=True)

##创建word
def Set_Background_Color(cell,rgbColor):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = rgbColor)) #固定写法，照抄即可
    cell._tc.get_or_add_tcPr().append(shading_elm)
def xlsx2docx(fn):
    #打开Excel文件，如果有公式的话，读取公式计算结果
    wb = load_workbook(fn)
    #创建空白Word文件
    document = Document()
    microsoft_font = u'微软雅黑'  # u 表示后面的字符串以 Unicode 格式进行编码
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal'].font.size = Pt(7.5)
    document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)

    # 设置标题样式
    black_font = u'微软雅黑'
    run = document.add_heading('', level=0).add_run('广西机顶盒软探针指标分析报告 -'+ todaytime1)
    run.font.name = black_font
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(area, black_font)
    #添加段落
    document.add_paragraph('互联网电视用户数'+str(data1)+'万'+'，其中已部署机顶盒软探针用户数'+str(data1)+'万'+'，部署比例100%；'+
                           '软探针日收视用户数'+str(data2)+'万'+'，活跃率'+str(data3)+'。')
    #设置一级标题
    black_font = u'微软雅黑'
    run = document.add_heading('', level=1).add_run('1.各地市日收视用户情况')
    run.font.name = black_font
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(area, black_font)
        #查看所有可用的表格样式
    table_styles = [style for style in document.styles
                    if style.type == WD_STYLE_TYPE.TABLE]

    #遍历Excel文件中所有的worksheet
    for WS in wb.worksheets:
        rows = list(WS.rows)
        #根据Worksheet的行数和列数，在Word文件中创建合适大小的表格
        table = document.add_table(rows=len(rows),
                                    cols=len(rows[0]),
                                    style='Table Grid')
    #从Worksheet读取数据，写入Word文件中的表格
        for irow,row in enumerate(rows):
            for icol,col in enumerate(row):
                table.cell(irow,icol).text = str(col.value)
                table.style.font.color.rgb = RGBColor(0, 0, 0)
    #保存Word文件
    document.save(fn[:-4]+'docx' )
#调用函数，进行数据导入
xlsx2docx('广西机顶盒软探针指标分析报告-'+todaytime3 + '.xlsx')

#
doc = Document('广西机顶盒软探针指标分析报告-'+todaytime3 + '.docx') #打开word文件
table1= doc.tables[1]
max_row = len(table1.rows) #获取表格总行数
#有线接入率
qty = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[1].text  # cells[x+1]指表格第x列
    qty_info = float(qty_info.strip('%'))
    qty_info = qty_info / 100
    qty.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty:
    if i > 0.92:
        if i > 0.96:
            cell = table1.cell(row, 1)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
            row += 1  # 跳转到下一行
        else:
            cell = table1.cell(row, 1)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFF00")  # 填充颜色黄色
            row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 1)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
# EPG响应时延（ms）
qty7 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[2].text  # cells[x+1]指表格第x列
    qty7.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty7:
    if i < 500:
        if i < 200:
            cell = table1.cell(row, 2)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
            row += 1  # 跳转到下一行
        else:
            cell = table1.cell(row, 2)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFF00")  # 填充颜色黄色
            row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 2)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
#EPG响应成功率
qty1 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[3].text  # cells[x+1]指表格第x列
    qty_info = float(qty_info.strip('%'))
    qty_info = qty_info / 100
    qty1.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty1:
    if i > 0.97:
        if i > 0.99:
            cell = table1.cell(row, 3)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
            row += 1  # 跳转到下一行
        else:
            cell = table1.cell(row, 3)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFF00")  # 填充颜色黄色
            row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 3)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
#播放成功率
qty2 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[4].text  # cells[x+1]指表格第x列
    qty_info = float(qty_info.strip('%'))
    qty_info = qty_info / 100
    qty2.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty2:
    if i > 0.99:
        if i > 0.995:
            cell = table1.cell(row, 4)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
            row += 1  # 跳转到下一行
        else:
            cell = table1.cell(row, 4)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFF00")  # 填充颜色黄色
            row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 4)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
#首次加载时长（ms）
qty3 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[5].text  # cells[x+1]指表格第x列
    qty3.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty3:
    if i < 2000:
        if i < 500:
            cell = table1.cell(row, 5)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
            row += 1  # 跳转到下一行
        else:
            cell = table1.cell(row, 5)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFF00")  # 填充颜色黄色
            row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 5)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
# 卡顿时长占比
qty4 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[6].text  # cells[x+1]指表格第x列
    qty_info = float(qty_info.strip('%'))
    qty_info = qty_info / 100
    qty4.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty4:
    if i < 0.0007:
        if i < 0.0005:
            cell = table1.cell(row, 6)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
            row += 1  # 跳转到下一行
        else:
            cell = table1.cell(row, 6)  # 第列数据为数量，列索引是
            Set_Background_Color(cell, "FFFF00")  # 填充颜色黄色
            row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 6)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
# 卡顿用户占比
qty5 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[7].text  # cells[x+1]指表格第x列
    qty_info = float(qty_info.strip('%'))
    qty_info = qty_info / 100
    qty5.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty5:
    if i < 0.01:
        cell = table1.cell(row, 7)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FFFFFF")  # 填充颜色白色
        row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 7)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
# 视频播放优良率d-OA
qty6 = []  # 存储数量信息
# 读取第二行到末尾行自选的数据
for i in range(1, max_row):
    qty_info = table1.rows[i].cells[8].text  # cells[x+1]指表格第x列
    qty_info = float(qty_info.strip('%'))
    qty_info = qty_info / 100
    qty6.append(float(qty_info))
# 判断
row = 1  # 行计数器
for i in qty6:
    if i < 0.97:
        cell = table1.cell(row, 8)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FF0000")  # 填充颜色红色
        row += 1  # 跳转到下一行
    else:
        cell = table1.cell(row, 8)  # 第列数据为数量，列索引是
        Set_Background_Color(cell, "FFFFFF")  # 填充颜色红色
        row += 1  # 跳转到下一行
area = qn('w:eastAsia')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.各地市关键指标情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#设置二级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.1各地市关键指标地市间对比情况：')
run.font.name = black_font
run.font.bold = False
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('注1：表格中标注红色底色的指标是未达到基准值，标注黄色底色的指标是已达到基准值未达到挑战值，无颜色标注的指标是已达到挑战值。')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2 各地市关键指标异常情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.1 各地市有线接入率未达标情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('(有线接入率——基准值：92%，挑战值：96%)')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.1 各地市EPG响应时延（ms）未达标情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('（EPG响应时延（ms）——基准值：500ms，挑战值：200ms）')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.2 各地市EPG响应成功率（%）指标未达标情况:')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('(EPG响应成功率（%）——基准值：97%，挑战值：99%)')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.3 各地市播放成功率（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('（播放成功率（%）——基准值：99%，挑战值：99.5%）')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.4 各地市平均首次加载时长（s）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('（平均首次加载时长（s）——基准值：2s，挑战值：0.5s）')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.5 各地市卡顿/花屏时长占比（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('（卡顿/花屏时长占比（%）——基准值：0.07%，挑战值：0.05%）')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.6 各地市卡顿/花屏用户占比（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('（卡顿/花屏用户占比（%）——基准值：1%）')
doc.add_paragraph('')
#设置一级标题
black_font = u'微软雅黑'
run = doc.add_heading('', level=1).add_run('2.2.7 各地市视频播放优良率d-OA占比（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
#添加段落
doc.add_paragraph('（视频播放优良率d-OA占比（%）——挑战值：97%）')
doc.add_paragraph('')
doc.save('广西机顶盒软探针指标分析报告-'+todaytime3 + '.docx')
# #获取表格文本，以列表形式返回
# def get_table_text(path, n=0):
#     """
#     获取word中的第n个表格的文本
#     path: word路径
#     n: 第几个表格，从0开始计算
#     :return: list类型的二维数组
#     """
#     document = Document(path)
#     all_tables = len(document.tables)
#     if all_tables > n:
#         table = document.tables[n]
#         text_list = []
#         for col in table.columns:
#             text = []
#             for cell in col.cells:
#                 text.append(cell.text)
#             text_list.append(text)
#
#         return text_list
#     else:
#         raise IndexError('table index (%s) out of range, in total %s' % (n, all_tables))
# a = get_table_text('111.docx', n=1)
# print(a)

# c = pd.DataFrame(a)
# df = pd.DataFrame(c.values.T, index=c.columns, columns=c.index)#转置
# list0 = list(df.iloc[0])
# df.columns = list0  ###重命名表头
# df = df.drop(df.head(1).index)
# df=df.reset_index(drop=True)
# print(df)





