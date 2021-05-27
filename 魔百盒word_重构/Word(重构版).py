import pandas as pd
import datetime
from docx import Document
from openpyxl import load_workbook
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

todaytime = datetime.datetime.now()  # 系统取当天日期
yesterday = todaytime - datetime.timedelta(days=1)  # 获取前一天的日期
yesterday = yesterday.strftime('%Y') + '-' + str(yesterday.month) + '-' + yesterday.strftime('%d')
####################################
def title(data,level=1):
    black_font = u'微软雅黑'
    run = document.add_heading('', level=level).add_run(data)
    run.font.name = black_font
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(area, black_font)
    return


def set_background_color(cell, rgbcolor):
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbcolor))  # 固定写法，照抄即可
    cell._tc.get_or_add_tcPr().append(shading_elm)

def write(x,y) :
    g = []
    if x != g:
        if y != g:
            document.add_paragraph(','.join(x) + '未达到基准值' + '，' + ','.join(y) + '达到了基准值,但未达挑战值')
        else:
            document.add_paragraph(','.join(x) + '未达到基准值')
    else:
        if y != g:
            document.add_paragraph(','.join(y) + '达到了基准值,但未达挑战值')
        else:
            document.add_paragraph('暂无')

w = pd.read_excel('基础数据整理-0423.xlsx', skiprows=38)  # 从38行开始读表
biao1 = w.iloc[:15, :9]  # 切片
biao2 = w.iloc[:15, 10:15]

# 转换百分比
biao1['有线接入率'] = biao1['有线接入率'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['EPG响应成功率'] = biao1['EPG响应成功率'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['播放成功率'] = biao1['播放成功率'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['卡顿时长占比'] = biao1['卡顿时长占比'].apply(lambda x: '%.4f%%' % (x * 100))
biao1['卡顿用户占比'] = biao1['卡顿用户占比'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['视频播放优良率d-OA'] = biao1['视频播放优良率d-OA'].apply(lambda x: '%.2f%%' % (x * 100))
biao2['收视率'] = biao2['收视率'].apply(lambda x: '%.2f%%' % (x * 100))
# 增加日期列
biao2.insert(1, '日期', yesterday)

with pd.ExcelWriter('基础数据整理.xlsx') as writer:  # 写入结果为当前路径
    biao2.to_excel(writer, sheet_name='表1', startcol=0, index=False, header=True)
    biao1.to_excel(writer, sheet_name='表2', startcol=0, index=False, header=True)

###########################################
date = str(biao2.iloc[14]['日期'])  # date
k1 = str(biao2.iloc[14]['上月用户数'])  # 空2
k2 = str(biao2.iloc[14]['上月用户数'])  # 空3
k3 = str(biao2.iloc[14]['日收视用户数'])  # 空4
k4 = str(biao2.iloc[14]['收视率'])  # 空5

#############################################################
# 打开Excel文件，如果有公式的话，读取公式计算结果
wb = load_workbook('基础数据整理.xlsx')
# 创建空白Word文件
document = Document()
g = []
# 设置字体
microsoft_font = u'微软雅黑'  # u 表示后面的字符串以 Unicode 格式进行编码
area = qn('w:eastAsia')
document.styles['Normal'].font.name = microsoft_font
document.styles['Normal'].font.size = Pt(7.5)
document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)

# 设置标题样式
black_font = u'微软雅黑'
run = document.add_heading('', level=0).add_run('广西机顶盒软探针指标分析报告 -' + yesterday)
run.font.name = black_font
run.font.bold = True #是否加粗
run.font.size = Pt(14) #设置字体大小14磅
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)

# 添加段落
document.add_paragraph('互联网电视用户数'+str(k1)+'万'+'，其中已部署机顶盒软探针用户数'+str(k1)+'万'+'，部署比例100%；' +
                           '软探针日收视用户数'+str(k3)+'万'+'，活跃率'+str(k4)+'。')
# 设置一级标题
title('1.各地市日收视用户情况')
# 遍历Excel文件中所有的worksheet
for WS in wb.worksheets:
    rows = list(WS.rows)
        # 根据Worksheet的行数和列数，在Word文件中创建合适大小的表格
    table = document.add_table(rows=len(rows),
                                    cols=len(rows[0]),
                                    style='Table Grid')
# 从Worksheet读取数据，写入Word文件中的表格
    for irow, row in enumerate(rows):
        for icol, col in enumerate(row):
            table.cell(irow, icol).text = str(col.value)
            table.style.font.color.rgb = RGBColor(0, 0, 0)
    document.add_paragraph('注1：日收视活跃率=日收视用户数/上月用户数*100%。')
# 输出判断
table1 = document.tables[1]
max_row = len(table1.rows)  # 获取表格总行数
    # 设置一级标题
title('2.各地市关键指标情况')
    # 设置二级标题
title('2.1各地市关键指标地市间对比情况：')
# 添加段落
document.add_paragraph('注1：表格中标注红色底色的指标是未达到基准值，标注黄色底色的指标是已达到基准值未达到挑战值，无颜色标注的指标是已达到挑战值。')
# 设置一级标题
title('2.2 各地市关键指标异常情况：')
# 设置一级标题
title('2.2.1 各地市有线接入率未达标情况')
# 添加段落
document.add_paragraph('(有线接入率——基准值：92%，挑战值：96%)')

# 有线接入率——基准值：92%，挑战值：96%
qty0 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty0_info = table1.rows[i].cells[1].text  # cells[1]指表格第2列
    qty0_info = float(qty0_info.strip('%'))
    qty0_info = qty0_info / 100
    qty0.append(float(qty0_info))
# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x = []
y = []
for i in qty0:
    if i >= 0.92:
        if i <= 0.96:
            cell = table1.cell(row, 1)  # 第2列数据为数量，列索引是1
            set_background_color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
            x.append(table1.cell(row-1, 0).text)
    else:
        cell = table1.cell(row, 1)  # 第2列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        y.append(table1.cell(row-1, 0).text)

    row += 1  # 跳转到下一行
write(x,y)
# 设置一级标题
title('2.2.1 各地市EPG响应时延（ms）未达标情况')
# 添加段落
document.add_paragraph('（EPG响应时延（ms）——基准值：500ms，挑战值：200ms）')
# EPG响应时延（ms）——基准值：500ms，挑战值：200ms
qty1 = []  # 存储数量信息
# 读取第二行到15行，第3列中的数据
for i in range(1, max_row):
    qty1_info = table1.rows[i].cells[2].text  # cells[2]指表格第3列
    qty1.append(float(qty1_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x1 = []
y1 = []
for i in qty1:
    if i < 500:
        if i > 200:
            cell = table1.cell(row, 2)  # 第2列数据为数量，列索引是1
            set_background_color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
            x1.append(table1.cell(row-1, 0).text)
    else:
        cell = table1.cell(row, 2)  # 第2列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        y1.append(table1.cell(row-1, 0).text)
    row += 1  # 跳转到下一行

write(x1,y1)

# 设置一级标题
title('2.2.2 各地市EPG响应成功率（%）指标未达标情况:')
# 添加段落
document.add_paragraph('(EPG响应成功率（%）——基准值：97%，挑战值：99%)')

qty2 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty2_info = table1.rows[i].cells[3].text  # cells[1]指表格第2列
    qty2_info = float(qty2_info.strip('%'))
    qty2_info = qty2_info / 100
    qty2.append(float(qty2_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x2 = []
y2 = []
for i in qty2:
    if i >= 0.97:
        if i <= 0.99:
            cell = table1.cell(row, 3)  # 第2列数据为数量，列索引是1
            set_background_color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
            x2.append(table1.cell(row-1, 0).text)
    else:
        cell = table1.cell(row, 3)  # 第2列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        y2.append(table1.cell(row-1, 0).text)

    row += 1  # 跳转到下一行

write(x2,y2)

# 设置一级标题
title('2.2.3 各地市播放成功率（%）指标未达标情况：')

# 添加段落
document.add_paragraph('（播放成功率（%）——基准值：99%，挑战值：99.5%）')
qty3 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty3_info = table1.rows[i].cells[4].text  # cells[1]指表格第2列
    qty3_info = float(qty3_info.strip('%'))
    qty3_info = qty3_info / 100
    qty3.append(float(qty3_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x3 = []
y3 = []
for i in qty3:
    if i > 0.99:
        if i <= 0.995:
            cell = table1.cell(row, 4)  # 第2列数据为数量，列索引是1
            set_background_color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
            x3.append(table1.cell(row-1, 0).text)
    else:
        cell = table1.cell(row, 4)  # 第2列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        y3.append(table1.cell(row-1, 0).text)

    row += 1  # 跳转到下一行

write(x3,y3)
# 设置一级标题
title('2.2.4 各地市平均首次加载时长（s）指标未达标情况：')

# 添加段落
document.add_paragraph('（平均首次加载时长（s）——基准值：2s，挑战值：0.5s）')
qty4 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty4_info = table1.rows[i].cells[5].text  # cells[1]指表格第2列
    qty4.append(float(qty4_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x4 = []
y4 = []
for i in qty4:
    if i <= 2000:
        if i >= 500:
            cell = table1.cell(row, 5)  # 第2列数据为数量，列索引是1
            set_background_color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
            x4.append(table1.cell(row-1, 0).text)
    else:
        cell = table1.cell(row, 5)  # 第2列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        y4.append(table1.cell(row-1, 0).text)

    row += 1  # 跳转到下一行

write(x4,y4)

# 设置一级标题
title('2.2.5 各地市卡顿/花屏时长占比（%）指标未达标情况：')

# 添加段落
document.add_paragraph('（卡顿/花屏时长占比（%）——基准值：0.07%，挑战值：0.05%）')
qty5 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty5_info = table1.rows[i].cells[6].text  # cells[1]指表格第2列
    qty5_info = float(qty5_info.strip('%'))
    qty5_info = qty5_info / 100
    qty5.append(float(qty5_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x5 = []
y5 = []
for i in qty5:
    if i < 0.0007:
        if i > 0.0005:
            cell = table1.cell(row, 6)  # 第2列数据为数量，列索引是1
            set_background_color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
            x5.append(table1.cell(row-1, 0).text)
    else:
        cell = table1.cell(row, 6)  # 第2列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        y5.append(table1.cell(row-1, 0).text)

    row += 1  # 跳转到下一行

write(x5,y5)

# 设置一级标题
title('2.2.6 各地市卡顿/花屏用户占比（%）指标未达标情况：')
# 添加段落
document.add_paragraph('（卡顿/花屏用户占比（%）——基准值：1%）')
qty6 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty6_info = table1.rows[i].cells[7].text  # cells[1]指表格第2列
    qty6_info = float(qty6_info.strip('%'))
    qty6_info = qty6_info / 100
    qty6.append(float(qty6_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x6 = []
for i in qty6:
    if i >= 0.01:
        cell = table1.cell(row, 7)  # 第7列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        x6.append(table1.cell(row-1, 0).text)
    row += 1  # 跳转到下一行

if x6 != g:
    document.add_paragraph(','.join(x6) + '未达到基准值')
else:
    document.add_paragraph('暂无')

# 设置一级标题
title('2.2.7 各地市视频播放优良率d-OA占比（%）指标未达标情况：')
# 添加段落
document.add_paragraph('（视频播放优良率d-OA占比（%）——挑战值：97%）')
qty7 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty7_info = table1.rows[i].cells[8].text  # cells[1]指表格第2列
    qty7_info = float(qty7_info.strip('%'))
    qty7_info = qty7_info / 100
    qty7.append(float(qty7_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
x7 = []
for i in qty7:
    if i <= 0.97:
        cell = table1.cell(row, 8)  # 第7列数据为数量，列索引是1
        set_background_color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码
        x7.append(table1.cell(row-1, 0).text)
    row += 1  # 跳转到下一行

if x7 != g:
    document.add_paragraph(','.join(x7) + '未达到挑战值')
else:
    document.add_paragraph('暂无')

document.save("广西机顶盒软探针指标分析报告.docx")