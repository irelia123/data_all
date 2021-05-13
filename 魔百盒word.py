import pandas as pd
import datetime
from docx import Document
from openpyxl import load_workbook
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor,Pt
from docx.oxml.ns import qn
import xlrd


####################################
colorStr = '#FFFF00'


def tabBgColor(table, row, cols, colorStr):
    # print(table)
    # for i , vl in enumerate(table):
    #     print(vl)
    shading_list = locals()
    # for i in range(cols):
    shading_list['shading_elm_' + str(cols)] = parse_xml(
        r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))  # color
    table.rows[row].cells[cols]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(cols)])


def tabif(col, value):
    if col == 1:
        if value == "92%":
            print()
    print()


w = pd.read_excel('基础数据整理-0423.xlsx', skiprows=38)  # 从38行开始读表
biao1 = w.iloc[:15, :9]  # 切片
biao2 = w.iloc[:15, 10:15]

########################################转换百分比
biao1['有线接入率'] = biao1['有线接入率'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['EPG响应成功率'] = biao1['EPG响应成功率'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['播放成功率'] = biao1['播放成功率'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['卡顿时长占比'] = biao1['卡顿时长占比'].apply(lambda x: '%.4f%%' % (x * 100))
biao1['卡顿用户占比'] = biao1['卡顿用户占比'].apply(lambda x: '%.2f%%' % (x * 100))
biao1['视频播放优良率d-OA'] = biao1['视频播放优良率d-OA'].apply(lambda x: '%.2f%%' % (x * 100))
biao2['收视率'] = biao2['收视率'].apply(lambda x: '%.2f%%' % (x * 100))

##########################增加日期列
todaytime = datetime.datetime.now()  # 系统取当天日期
yesterday = todaytime - datetime.timedelta(days=1)  # 获取前一天的日期
yesterday = yesterday.strftime('%Y') + '-' + str(yesterday.month) + '-' + yesterday.strftime('%d')
biao2.insert(1, '日期', yesterday)
todaytime3 = todaytime.strftime('%m')+'-'+ todaytime.strftime('%d')

with pd.ExcelWriter('基础数据整理.xlsx') as writer:  # 写入结果为当前路径
    biao2.to_excel(writer, sheet_name='表1', startcol=0, index=False, header=True)
    biao1.to_excel(writer, sheet_name='表2', startcol=0, index=False, header=True)

###########################################
date = str(biao2.iloc[14]['日期'])  # date
k1 = str(biao2.iloc[14]['上月用户数'])  # 空2
k2 = str(biao2.iloc[14]['上月用户数'])  # 空3
k3 = str(biao2.iloc[14]['日收视用户数'])  # 空4
k4 = str(biao2.iloc[14]['收视率'])  # 空5

doc = Document("广西机顶盒软探针指标分析报告-0423.docx")
data = [date, k1, k2, k3, k4]

c = 0
for p in doc.paragraphs:
    if 'XXXX' in p.text:
        inline = p.runs
        for i in range(len(inline)):
            if "XXXX" in inline[i].text:
                text = inline[i].text.replace('XXXX', data[c])
                inline[i].text = text
                c = c + 1



#############################################################
#def xlsx2docx(fn):
    #打开Excel文件，如果有公式的话，读取公式计算结果
wb = load_workbook('基础数据整理.xlsx')
    #创建空白Word文件
document = Document()
    # 设置颜色
rde = RGBColor(255, 0, 0)
yellow = RGBColor(255, 255, 0)
    #设置字体
microsoft_font = u'微软雅黑'  # u 表示后面的字符串以 Unicode 格式进行编码
area = qn('w:eastAsia')
document.styles['Normal'].font.name = microsoft_font
document.styles['Normal'].font.size = Pt(7.5)
document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)

    # 设置标题样式
black_font = u'微软雅黑'
run = document.add_heading('', level=0).add_run('广西机顶盒软探针指标分析报告 -'+ yesterday)
run.font.name = black_font
run.font.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('互联网电视用户数'+str(k1)+'万'+'，其中已部署机顶盒软探针用户数'+str(k1)+'万'+'，部署比例100%；'+
                           '软探针日收视用户数'+str(k3)+'万'+'，活跃率'+str(k4)+'。')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('1.各地市日收视用户情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
        #查看所有可用的表格样式
    # table_styles = [style for style in document.styles
    #                 if style.type == WD_STYLE_TYPE.TABLE]

    #遍历Excel文件中所有的worksheet
for WS in wb.worksheets:
    rows = list(WS.rows)
        #根据Worksheet的行数和列数，在Word文件中创建合适大小的表格
    table = document.add_table(rows=len(rows),
                                    cols=len(rows[0]),
                                    style='Table Grid')
    #从Worksheet读取数据，写入Word文件中的表格
    for irow,row in enumerate(rows):
        for icol,col in enumerate(row):
            table.cell(irow,icol).text = str(col.value)
            table.style.font.color.rgb = RGBColor(0, 0, 0)
    document.add_paragraph('注1：日收视活跃率=日收视用户数/上月用户数*100%。')

    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.各地市关键指标情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #设置二级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.1各地市关键指标地市间对比情况：')
run.font.name = black_font
run.font.bold = False
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('注1：表格中标注红色底色的指标是未达到基准值，标注黄色底色的指标是已达到基准值未达到挑战值，无颜色标注的指标是已达到挑战值。')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2 各地市关键指标异常情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.1 各地市有线接入率未达标情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('(有线接入率——基准值：92%，挑战值：96%)')

#输出判断
path = '基础数据整理.xlsx'  # 源数据的路径
wb1 = xlrd.open_workbook(path)
sh = wb1.sheet_by_name("表2")
nrows = sh.nrows  # 获取行数
x = []  #未达到基准值
y = []  #达到了基准值,但未达挑战值
z = []  #达到挑战值
p = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 1) < '92%':
        a = sh.cell_value(i, 0)  #未达到基准值
        x.append(a) #将地市放到列表
    elif sh.cell_value(i, 1) >= '92%' and sh.cell_value(i, 1) < '96%':
        b = sh.cell_value(i, 0)  #达到了基准值,但未达挑战值
        y.append(b)
    elif sh.cell_value(i, 1) >= '96%':
        c = sh.cell_value(i, 0) #达到挑战值
        z.append(c)
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p.append(d)

g = []
if x != g:
    if y != g:
        document.add_paragraph(','.join(x) + '未达到基准值' + '，' + ','.join(y) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph(','.join(x) + '未达到基准值')
else:
    if y != g:
        document.add_paragraph(','.join(y) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph('暂无')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.1 各地市EPG响应时延（ms）未达标情况')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('（EPG响应时延（ms）——基准值：500ms，挑战值：200ms）')
    #设置一级标题
x1 = []  #未达到基准值
y1 = []  #达到了基准值,但未达挑战值
z1 = []  #达到挑战值
p1 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 2) > 500 :
        a = sh.cell_value(i, 0)  #未达到基准值
        x1.append(a) #将地市放到列表
    elif sh.cell_value(i, 2) > 200 and sh.cell_value(i, 2) <= 500 :
        b = sh.cell_value(i, 0)  #达到了基准值,但未达挑战值
        y1.append(b)
    elif sh.cell_value(i, 2) <= 200:
        c = sh.cell_value(i, 0) #达到挑战值
        z1.append(c)
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p1.append(d)

g1 = []
if x != g1:
    if y != g1:
        document.add_paragraph(','.join(x1) + '未达到基准值' + '，' + ','.join(y1) +'达到了基准值,但未达挑战值')
    else:
        document.add_paragraph(','.join(x1)+'未达到基准值')
else:
    if y1 != g1:
        document.add_paragraph(','.join(y1) +'达到了基准值,但未达挑战值')
    else:
        document.add_paragraph('暂无')



black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.2 各地市EPG响应成功率（%）指标未达标情况:')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('(EPG响应成功率（%）——基准值：97%，挑战值：99%)')
#
x2 = []  #未达到基准值
y2 = []  #达到了基准值,但未达挑战值
z2 = []  #达到挑战值
p2 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 3) < '97%':
        a = sh.cell_value(i, 0)  #未达到基准值
        x2.append(a) #将地市放到列表
    elif sh.cell_value(i, 3) >= '97%' and sh.cell_value(i, 3) < '99%':
        b = sh.cell_value(i, 0)  #达到了基准值,但未达挑战值
        y2.append(b)
    elif sh.cell_value(i, 3) >= '99%':
        c = sh.cell_value(i, 0) #达到挑战值
        z2.append(c)
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p2.append(d)

g2 = []
if x2 != g2:
    if y2 != g2:
        document.add_paragraph(','.join(x2) + '未达到基准值' + '，' + ','.join(y2) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph(','.join(x2) + '未达到基准值')
else:
    if y2 != g2:
        document.add_paragraph(','.join(y2) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph('暂无')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.3 各地市播放成功率（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('（播放成功率（%）——基准值：99%，挑战值：99.5%）')
#
x3 = []  #未达到基准值
y3 = []  #达到了基准值,但未达挑战值
z3 = []  #达到挑战值
p3 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 4) < '99%':
        a = sh.cell_value(i, 0)  #未达到基准值
        x3.append(a) #将地市放到列表
    elif sh.cell_value(i, 4) >= '99%' and sh.cell_value(i, 4) <'99.50%':
        b = sh.cell_value(i, 0)  #达到了基准值,但未达挑战值
        y3.append(b)
    elif sh.cell_value(i, 4) >= '99.50%':
        c = sh.cell_value(i, 0) #达到挑战值
        z3.append(c)
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p3.append(d)

g3 = []
if x3 != g3:
    if y3 != g3:
        document.add_paragraph(','.join(x3) + '未达到基准值' + '，' + ','.join(y3) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph(','.join(x3) + '未达到基准值')
else:
    if y3 != g3:
        document.add_paragraph(','.join(y3) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph('暂无')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.4 各地市平均首次加载时长（s）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('（平均首次加载时长（s）——基准值：2s，挑战值：0.5s）')
x4 = []  #未达到基准值
y4 = []  #达到了基准值,但未达挑战值
z4 = []  #达到挑战值
p4 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 5) > 2000:
        a = sh.cell_value(i, 0)  #未达到基准值
        x4.append(a) #将地市放到列表
    elif sh.cell_value(i, 5) >500  and sh.cell_value(i, 5) <= 2000 :
        b = sh.cell_value(i, 0)  #达到了基准值,但未达挑战值
        y4.append(b)
    elif sh.cell_value(i, 5) <= 500 :
        c = sh.cell_value(i, 0) #达到挑战值
        z4.append(c)
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p4.append(d)

g4 = []
if x4 != g4:
    if y4 != g4:
        document.add_paragraph(','.join(x4) + '未达到基准值' + '，' + ','.join(y4) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph(','.join(x4) + '未达到基准值')
else:
    if y4 != g4:
        document.add_paragraph(','.join(y4) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph('暂无')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.5 各地市卡顿/花屏时长占比（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('（卡顿/花屏时长占比（%）——基准值：0.07%，挑战值：0.05%）')
x5 = []  #未达到基准值
y5 = []  #达到了基准值,但未达挑战值
z5 = []  #达到挑战值
p5 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 6) >= '0.0700%':
        a = sh.cell_value(i, 0)  #未达到基准值
        x5.append(a) #将地市放到列表
    elif sh.cell_value(i, 6) >= '0.0500%' and sh.cell_value(i, 6) < '0.0700%':
        b = sh.cell_value(i, 0)  #达到了基准值,但未达挑战值
        y5.append(b)
    elif sh.cell_value(i, 6) <= '0.0500%' :
        c = sh.cell_value(i, 0) #达到挑战值
        z5.append(c)
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p5.append(d)

g5 = []
if x5 != g5:
    if y5 != g5:
        document.add_paragraph(','.join(x5) + '未达到基准值' + '，' + ','.join(y5) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph(','.join(x5) + '未达到基准值')
else:
    if y5 != g5:
        document.add_paragraph(','.join(y5) + '达到了基准值,但未达挑战值')
    else:
        document.add_paragraph('暂无')
    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.6 各地市卡顿/花屏用户占比（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('（卡顿/花屏用户占比（%）——基准值：1%）')
x6 = []  #未达到基准值
y6 = []  #达到了基准值,但未达挑战值
z6 = []  #达到挑战值
p6 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 7) > '1%':
        a = sh.cell_value(i, 0)  #未达到基准值
        x6.append(a) #将地市放到列表
    else:
        d = sh.cell_value(i, 0)  #达到基准值
        p6.append(d)
g6 = []
if x6 != g6:

    document.add_paragraph(','.join(x6) + '未达到基准值')
else:
    document.add_paragraph('暂无')



    #设置一级标题
black_font = u'微软雅黑'
run = document.add_heading('', level=1).add_run('2.2.7 各地市视频播放优良率d-OA占比（%）指标未达标情况：')
run.font.name = black_font
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(area, black_font)
    #添加段落
document.add_paragraph('（视频播放优良率d-OA占比（%）——挑战值：97%）')
x7 = []  #未达到基准值
y7 = []  #达到了基准值,但未达挑战值
z7 = []  #达到挑战值
p7 = []  #达到基准值
for i in range(1,nrows):
    if sh.cell_value(i, 8) < '97%':
        a = sh.cell_value(i, 0)  #未达到挑战值
        x7.append(a) #将地市放到列表
    else:
        d = sh.cell_value(i, 0)  #达到挑战值
        p7.append(d)
g7 = []
if x7 != g7:

    document.add_paragraph(','.join(x7) + '未达到挑战值')
else:
    document.add_paragraph('暂无')


#保存Word文件
document.save('广西机顶盒软探针指标分析报告.docx' )
#调用函数，进行数据导入
#xlsx2docx('基础数据整理.xlsx')
#填充颜色
doc = Document("广西机顶盒软探针指标分析报告.docx") #打开word文件
table1 =doc.tables[1]
max_row = len(table1.rows) #获取表格总行数
def Set_Background_Color(cell, rgbColor):
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))  # 固定写法，照抄即可
    cell._tc.get_or_add_tcPr().append(shading_elm)


# 有线接入率——基准值：92%，挑战值：96%
qty0 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty0_info = table1.rows[i].cells[1].text  # cells[1]指表格第2列
    qty0_info = float(qty0_info.strip('%'))
    qty0_info = qty0_info / 100
    qty0.append(float(qty0_info))

# 将小于基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
for i in qty0:
    if i >= 0.92:
        if i <= 0.96:
            cell = table1.cell(row, 1)  # 第2列数据为数量，列索引是1
            Set_Background_Color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
    else:
        cell = table1.cell(row, 1)  # 第2列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行

# EPG响应时延（ms）——基准值：500ms，挑战值：200ms
qty1 = []  # 存储数量信息
# 读取第二行到15行，第3列中的数据
for i in range(1, max_row):
    qty1_info = table1.rows[i].cells[2].text  # cells[2]指表格第3列
    qty1.append(float(qty1_info))

# 将未达到基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
for i in qty1:
    if i < 500:
        if i > 200:
            cell = table1.cell(row, 2)  # 第3列数据为数量，列索引是1
            Set_Background_Color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
    else:
        cell = table1.cell(row, 2)  # 第3列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行

# EPG响应成功率（%）——基准值：97%，挑战值：99%
qty2 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty2_info = table1.rows[i].cells[3].text  # cells[1]指表格第4列
    qty2_info = float(qty2_info.strip('%'))
    qty2_info = qty2_info / 100
    qty2.append(float(qty2_info))

# 将未达到基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
for i in qty2:
    if i >= 0.97:
        if i <= 0.99:
            cell = table1.cell(row, 3)  # 第4列数据为数量，列索引是1
            Set_Background_Color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
    else:
        cell = table1.cell(row, 3)  # 第4列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行


# 播放成功率（%）——基准值：99%，挑战值：99.5%
qty3 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty3_info = table1.rows[i].cells[4].text  # cells[1]指表格第5列
    qty3_info = float(qty3_info.strip('%'))
    qty3_info = qty3_info / 100
    qty3.append(float(qty3_info))

# 将未达到基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
for i in qty3:
    if i > 0.99:
        if i <= 0.995:
            cell = table1.cell(row, 4)  # 第5列数据为数量，列索引是1
            Set_Background_Color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
    else:
        cell = table1.cell(row, 4)  # 第5列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行


# 平均首次加载时长（s）——基准值：2s，挑战值：0.5s
qty4 = []  # 存储数量信息
# 读取第二行到15行，第6列中的数据
for i in range(1, max_row):
    qty4_info = table1.rows[i].cells[5].text  # cells[1]指表格第6列
    qty4.append(float(qty4_info))

# 将未达到基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
for i in qty4:
    if i <= 2000:
        if i >= 500:
            cell = table1.cell(row, 5)  # 第6列数据为数量，列索引是1
            Set_Background_Color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
    else:
        cell = table1.cell(row, 5)  # 第6列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行


# 卡顿/花屏时长占比（%）——基准值：0.07%，挑战值：0.05%
qty5 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty5_info = table1.rows[i].cells[6].text  # cells[1]指表格第7列
    qty5_info = float(qty5_info.strip('%'))
    qty5_info = qty5_info / 100
    qty5.append(float(qty5_info))

# 将未达到基准值的单元格填红色，达到基准值但未达到挑战值的单元格填黄色
row = 1  # 行计数器
for i in qty5:
    if i < 0.0007:
        if i > 0.0005:
            cell = table1.cell(row, 6)  # 第7列数据为数量，列索引是1
            Set_Background_Color(cell, "#FFFF00")  # 填充颜色，"FFFF00"是黄色的编码
    else:
        cell = table1.cell(row, 6)  # 第7列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行

# 卡顿/花屏用户占比（%）——基准值：1%
qty6 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty6_info = table1.rows[i].cells[7].text  # cells[1]指表格第7列
    qty6_info = float(qty6_info.strip('%'))
    qty6_info = qty6_info / 100
    qty6.append(float(qty6_info))

# 将未达到基准值的单元格填红色
row = 1  # 行计数器
for i in qty6:
    if i >= 0.01:
        cell = table1.cell(row, 7)  # 第7列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

        row += 1  # 跳转到下一行

# 视频播放优良率d-OA占比（%）——挑战值：97%
qty7 = []  # 存储数量信息
# 读取第二行到15行，第2列中的数据
for i in range(1, max_row):
    qty7_info = table1.rows[i].cells[8].text  # cells[1]指表格第7列
    qty7_info = float(qty7_info.strip('%'))
    qty7_info = qty7_info / 100
    qty7.append(float(qty7_info))

# 将未达到基准值的单元格填红色
row = 1  # 行计数器
for i in qty7:
    if i <= 0.97:
        cell = table1.cell(row, 8)  # 第7列数据为数量，列索引是1
        Set_Background_Color(cell, "#FF0000")  # 填充颜色，"FF0000"是红色的编码

    row += 1  # 跳转到下一行

document.save('广西机顶盒软探针指标分析报告.docx' )
